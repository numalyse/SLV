import sys
import json
import os
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_table_fixed_layout(table):
    """
    Force la disposition fixe du tableau : Word respecte alors les largeurs
    de colonnes imposées au lieu de les élargir selon le contenu.
    """
    tblPr = table._tbl.tblPr

    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')

def format_time(time_val):
    """
    Formate le temps si le C++ envoie des millisecondes brutes, 
    ou renvoie le texte direct si le C++ a déjà formaté la chaîne.
    """
    if isinstance(time_val, str):
        return time_val
    
    seconds = time_val / 1000.0
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    f = int((seconds - int(seconds)) * 100)
    return f"{h:02d}:{m:02d}:{s:02d}:{f:02d}"

def main():
    if len(sys.argv) < 3:
        print("Erreur: Chemin d'un des fichiers JSON manquant.")
        sys.exit(1)

    data_json_path = sys.argv[1]
    #lang_json_path = sys.argv[2]
    shot_name = sys.argv[2]
    start_time_name = sys.argv[3]
    duration_time_name = sys.argv[4]
    image_label = sys.argv[5] if len(sys.argv) > 5 else "Image"
    sound_label = sys.argv[6] if len(sys.argv) > 6 else "Son"

    with open(data_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    temp_dir = data['tempDir']
    dst_path = data['dstPath']

    #with open(lang_json_path, 'r', encoding='utf-8') as f:
    #    lang = json.load(f)

    #shot_name = lang['shot']
    # start_time_name = lang['shot_detail_end_time_name']
    # duration_time_name = lang['shot_detail_duration_time_name']

    if not dst_path.endswith('.docx'):
        dst_path += '.docx'

    shots = data['shots']
    total_shots = len(shots)

    last_percent = -1

    try:
        doc = Document()
        
        titre_paragraphe = doc.add_paragraph()
        titre_paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        titre_run = titre_paragraphe.add_run(data["mediaName"].split(".")[0])
        titre_run.font.bold = True
        titre_run.font.size = Pt(24)
        titre_run.font.color.rgb = RGBColor.from_string("EB4034")

        doc.add_paragraph()

        for idx, shot in enumerate(shots):
            time_str = format_time(shot.get('start', 0))
            end_str = format_time(shot.get('duration', 0))
            plan_title = shot.get('title', f"{shot_name} {idx+1}")
            
            heading_text = f"- [{shot_name} {idx+1}] {plan_title} -> {start_time_name} : {time_str} / {duration_time_name} : {end_str}"
            heading = doc.add_heading(heading_text, level=2)
            
            run_heading = heading.runs[0]
            run_heading.font.color.rgb = RGBColor.from_string("1F497D")

            img_path = os.path.join(temp_dir, shot['image'])
            if os.path.exists(img_path):
                try:
                    with Image.open(img_path) as img:
                        px_width, px_height = img.size

                    img_docx_width = Inches(px_width / 96.0)
                    if img_docx_width > Inches(6.0):
                        img_docx_width = Inches(6.0)

                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    p_img.add_run().add_picture(img_path, width=img_docx_width)
                    
                except Exception as e:
                    print(f"Erreur lors du traitement de l'image {idx} : {e}", file=sys.stderr)

            img_text = shot.get('imgTxt', "").strip()
            sound_text = shot.get('soundTxt', "").strip()

            # On coupe la zone de note en deux colonnes : Image à gauche, Son à droite.
            # Largeur fixe pour que le texte long passe à la ligne DANS la cellule
            # au lieu d'élargir la colonne et de pousser la mise en page.
            column_width = Inches(3.0)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            _set_table_fixed_layout(table)
            for column in table.columns:
                column.width = column_width

            for cell, label, text in (
                (table.cell(0, 0), image_label, img_text),
                (table.cell(0, 1), sound_label, sound_text),
            ):
                cell.width = column_width
                cell_p = cell.paragraphs[0]
                label_run = cell_p.add_run(f"{label} : ")
                label_run.font.bold = True
                label_run.font.color.rgb = RGBColor.from_string("1F497D")
                cell_p.add_run(text)

            doc.add_paragraph()

            percent = int(((idx + 1) / total_shots) * 100)
            
            if percent != last_percent:
                print(f"PROGRESS:{percent}", flush=True)
                last_percent = percent 

        doc.save(dst_path)
        print("Fichier DOCX enregistré !", flush=True)

    except Exception as e:
        print(f"Erreur lors de l'exportation DOCX : {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()